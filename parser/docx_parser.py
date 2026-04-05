"""
parser/docx_parser.py — Rule-based, không LLM.
Cấu trúc đề cương NEU chuẩn (QĐ 628/2024).

Changelog v3 (bug-fix release):
  - _parse_instructors(): chấp nhận header "Họ tên" (không cần "và"), tìm email trong toàn bộ row
  - _parse_goals(): code_col nhận dạng "TT"/"STT"; fallback data-driven (scan G/CG codes)
  - _parse_resources(): chấp nhận trích dẫn plain text (không bắt buộc bullet/số)
  - _parse_general_info(): nhận dạng tín chỉ viết bằng chữ EN ("three-credit course")
  - _CREDITS_RE: thêm "units?" cho định dạng EN
  - parse_docx(): phát hiện is_thuc_te + has_clo_table → QA thông minh hơn
"""
import re
from io import BytesIO
from pathlib import Path
from typing import Optional
from docx import Document
from docx.table import Table
from .models import (
    CLO, Goal, Instructor, Resource, Assessment, Week,
    Requirement, Rubric, RubricCriterion, SyllabusRecord,
)

# ── Patterns ──────────────────────────────────────────────────────────────────
_VI_NAME_RE  = re.compile(r"tên\s*học\s*phần[^:]*tiếng\s*việt[^:]*[:\-]\s*(.+)", re.I)
_EN_NAME_RE  = re.compile(
    r"(?:tên\s*học\s*phần[^:]*tiếng\s*anh[^:]*[:\-]|title\s*:\s*|course\s*name\s*:\s*)\s*(.+)", re.I
)

_CODE_RE = re.compile(
    r"(?:mã\s*học\s*phần|(?<!\w)code|course\s*number)\s*[:\*\s]+"
    r"([A-Z]{2,}[0-9]{1,2}(?:\.[A-Z]{2,}[0-9]{2,})?|[A-Z0-9][A-Z0-9.\-]{3,})\b",
    re.I,
)
_CODE_FIELD_RE = re.compile(
    r"(?:mã\s*(?:hp|học\s*phần|môn)\s*[:\-]?\s*)"
    r"([A-Z]{2,}[0-9]{2,}(?:\.[A-Z]{2,}[0-9]{2,})?)",
    re.I,
)

# FIX: thêm "units?" để nhận dạng định dạng EN "- Units: 3"
# FIX: thêm "(\d+)TC" để nhận dạng "3TC" (KTKI1118 style)
_CREDITS_RE  = re.compile(
    r"(?:(?:số\s*tín\s*chỉ|credits?|units?)\s*[:\-\s]*\**\s*(\d+)\b"
    r"|(?<![.\d])(\d+)\s*TC(?!\w))",
    re.I,
)

# FIX MỚI: tín chỉ viết bằng chữ: "three-credit course", "two credits"
_WORD_CREDITS = {"one": 1, "two": 2, "three": 3, "four": 4, "five": 5, "six": 6}
_CREDITS_WORD_RE = re.compile(r"\b(one|two|three|four|five|six)\s*[-\u2013]?\s*credit", re.I)

_CLASS_H_RE  = re.compile(
    r"(?:số\s*giờ\s*trên\s*lớp|classroom\s*hours?|class\s*hours?|lecture\s*hours?|theoretical\s*hours?|lecture)\s*:\s*\**\s*(\d+)\b",
    re.I,
)
_SELF_H_RE   = re.compile(
    r"(?:số\s*giờ\s*tự\s*học|self.?study\s*hours?|self.?study)\s*[:\-]\s*\**\s*(\d+)\b",
    re.I,
)
_LEVEL_RE    = re.compile(r"(?:trình\s*độ\s*đào\s*tạo|level\s*of\s*training)\s*[:\-\s]*\**\s*(.+)", re.I)
_DECISION_RE = re.compile(
    r"(?:ban\s*hành\s*kèm\s*theo\s*quyết\s*định\s*số|issued\s*attached\s*decision)\s*[:\s]*(.+?)"
    r"(?:\s*ngày|\s*,|\s*\)|$)",
    re.I,
)
_PREREQ_HDR  = re.compile(r"(?:các\s*học\s*phần\s*tiên\s*quyết|(?:parallel|prerequisite)\s*course)", re.I)
_PREREQ_STP  = re.compile(
    r"(?:khoa[/\\]viện|địa\s*chỉ|giảng\s*viên|faculty:|address:|lecturer|"
    r"department|kind\s*of|meeting\s*hour|\+\s*faculty|\+\s*address)",
    re.I,
)
_FACULTY_RE  = re.compile(
    r"(?:khoa[/\\]viện\s*quản\s*lý\s*học\s*phần|department\s*conducting[^:]*)?\s*[:\+]\s*(.+)", re.I
)
_ADDRESS_RE  = re.compile(r"(?:địa\s*chỉ|address)\s*[:\+]\s*(.+)", re.I)
_EMAIL_RE    = re.compile(r"[\w.\-+]+@[\w.\-]+\.\w{2,}")

_SEC_DESC_RE   = re.compile(r"(?:\d+\.\s*)?(?:mô\s*tả\s*học\s*phần|course\s*descriptions?)\b", re.I)
_SEC_RES_RE    = re.compile(
    r"(?:\d+\.\s*)?(?:"
    r"tài\s*liệu\s*(?:học\s*tập|tham\s*khảo)"
    r"|learning\s*resources?"
    r"|required\s*textbooks?"
    r"|course\s*materials?"
    r"|recommended\s*texts?"
    r"|other\s*readings?"
    r")\b",
    re.I,
)
_SEC_GOAL_RE   = re.compile(r"(?:\d+\.\s*)?(?:mục\s*tiêu\s*học\s*phần|course\s*goals?)\b", re.I)
_SEC_CLO_RE    = re.compile(
    r"chuẩn\s*đầu\s*ra|course\s*learning\s*outcomes?|student\s*learning\s*outcomes?", re.I
)
_SEC_ASSESS_RE = re.compile(r"(?:\d+\.\s*)?(?:đánh\s*giá\s*học\s*phần|course\s*assessment)\b", re.I)
_SEC_WEEK_RE   = re.compile(r"(?:\d+\.\s*)?(?:kế\s*hoạch\s*dạy\s*học|lesson\s*plan)\b", re.I)
_SEC_REQ_RE    = re.compile(
    r"(?:\d+\.\s*)?(?:quy\s*định\s*học\s*phần|course\s*requirements?|quy\s*định\s*và\s*hướng\s*dẫn)\b",
    re.I,
)
_SEC_RUBRIC_RE = re.compile(r"(?:rubric|ma\s*trận\s*đề\s*thi|phụ\s*lục)", re.I)

_CLO_CODE_RE = re.compile(r"\bCLO\s*(\d+)(?:[.\-](\d+))?\b", re.I)
_GOAL_RE     = re.compile(r"\b(CG\d+|G\d+)\b", re.I)
_PLACEHOLDER = re.compile(r"^\s*\[\d+\]\s*$")

_ASSESSMENT_KIND_MAP = [
    (re.compile(r"chuyên\s*cần|attendance|participation", re.I), "attendance"),
    (re.compile(r"cuối\s*kỳ|final|thi\s*kết\s*thúc",    re.I), "final"),
    (re.compile(r"giữa\s*kỳ|midterm",                    re.I), "midterm"),
    (re.compile(r"nhóm|group",                            re.I), "group_project"),
    (re.compile(r"bài\s*tập|assignment|homework",         re.I), "assignment"),
    (re.compile(r"kiểm\s*tra|quiz|test",                  re.I), "quiz"),
]

_REQ_GROUP_MAP = [
    (re.compile(r"thi|kiểm\s*tra\s*cuối|điều\s*kiện\s*dự\s*thi|exam", re.I), "exam_condition"),
    (re.compile(r"chuyên\s*cần|điểm\s*danh|attendance",               re.I), "attendance"),
    (re.compile(r"hành\s*vi|thái\s*độ|trên\s*lớp|class",              re.I), "class_behavior"),
    (re.compile(r"nộp\s*bài|bài\s*tập|submission|assignment",         re.I), "assignment_submission"),
]


# ── Helpers ───────────────────────────────────────────────────────────────────

def _clean(t: str) -> str:
    return re.sub(r"\s+", " ", (t or "").replace("\u200b", "").replace("\xa0", " ")).strip()

def _is_placeholder_row(cells: list) -> bool:
    non_empty = [c for c in cells if c]
    return bool(non_empty) and all(_PLACEHOLDER.match(c) for c in non_empty)

def _norm_clo(major, minor) -> str:
    return f"CLO{major}.{minor}" if minor else f"CLO{major}"

def _roman_to_float(s: str) -> Optional[float]:
    rom = {"I": 1, "II": 2, "III": 3, "IV": 4, "V": 5, "VI": 6}
    s = s.strip()
    if s in rom: return float(rom[s])
    try:
        v = float(s.replace(",", "."))
        return v if 1.0 <= v <= 6.0 else None
    except ValueError:
        return None

def _extract_clo_codes(text: str) -> list:
    refs = []
    seen = set()
    for m in re.finditer(r"CLO\s*(\d+)[.\-](\d+)", text, re.I):
        code = f"CLO{m.group(1)}.{m.group(2)}"
        if code not in seen: refs.append(code); seen.add(code)
    for m in re.finditer(r"\b(\d+)\.(\d+)\s*[-–]\s*\1\.(\d+)\b", text):
        major, lo, hi = int(m.group(1)), int(m.group(2)), int(m.group(3))
        for minor in range(lo, hi + 1):
            code = f"CLO{major}.{minor}"
            if code not in seen: refs.append(code); seen.add(code)
    if not refs:
        for m in re.finditer(r"\b(\d+)\.(\d+)\b", text):
            code = f"CLO{m.group(1)}.{m.group(2)}"
            if code not in seen: refs.append(code); seen.add(code)
    return refs

def _guess_assessment_kind(title: str) -> str:
    for pattern, kind in _ASSESSMENT_KIND_MAP:
        if pattern.search(title):
            return kind
    return "other"

def _guess_req_group(text: str) -> str:
    for pattern, group in _REQ_GROUP_MAP:
        if pattern.search(text):
            return group
    return "other"

def _iter_para(doc: Document):
    for p in doc.paragraphs:
        t = _clean(p.text)
        if t: yield p, t

def _all_lines(doc: Document) -> list:
    lines = []
    for p in doc.paragraphs:
        t = _clean(p.text)
        if t: lines.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                t = _clean(cell.text)
                if t: lines.append(t)
    return lines


# ── Tên file → fallback code ──────────────────────────────────────────────────

def parse_filename(filename: str) -> tuple:
    stem  = Path(filename).stem
    parts = stem.split("_")

    # Dạng: EPxx.XYZ-nn (e.g. EP26.ETS-26) — phần sau dấu _ chứa dấu chấm và gạch ngang
    # Thử match toàn bộ đoạn cuối (sau dấu _) như là một code dạng PREFIX.SUFFIX
    COMPOUND_CODE = re.compile(
        r"^([A-Z]{2,}\d{1,2})\.([A-Z]{2,}[0-9\-]{2,})$", re.I
    )
    SINGLE_CODE = re.compile(
        r"^([A-Z]{2,}\d{1,2}\.[A-Z]{2,}\d{2,}|[A-Z]{2,}\d{2,}[A-Z0-9]*)$", re.I
    )
    PREFIX_PART  = re.compile(r"^(EP|MFE|FIN|ACC|MGT|MKT|IT|CS|SE|IS|DS)\d{1,2}$", re.I)
    CODE_PART    = re.compile(r"^[A-Z]{2,}\d{2,}$", re.I)

    if len(parts) >= 2 and PREFIX_PART.match(parts[-2]) and CODE_PART.match(parts[-1]):
        code = f"{parts[-2].upper()}.{parts[-1].upper()}"
        return code, " ".join(parts[:-2])

    if SINGLE_CODE.match(parts[-1]):
        return parts[-1].upper(), " ".join(parts[:-1])

    # FIX: strip trailing "(N)" suffix — e.g. "PTKT1102E (1)" → "PTKT1102E"
    last_cleaned = re.sub(r"\s*\(\d+\)\s*$", "", parts[-1]).strip()
    if last_cleaned != parts[-1] and SINGLE_CODE.match(last_cleaned):
        return last_cleaned.upper(), " ".join(parts[:-1])

    # FIX: dạng "EP26.ETS-26" — compound code với dấu gạch ngang ở phần suffix
    if COMPOUND_CODE.match(parts[-1]):
        return parts[-1].upper(), " ".join(parts[:-1])

    SEG = re.compile(r"^([A-Z]{2,}[0-9]{2,}|EP\d+)$", re.I)
    code_parts, i = [], len(parts) - 1
    while i >= 0 and SEG.match(parts[i]):
        code_parts.insert(0, parts[i]); i -= 1
    return ".".join(code_parts).upper() if code_parts else "", " ".join(parts[: i + 1])


# ── Section 1 & 2: thông tin cơ bản ──────────────────────────────────────────

def _parse_general_info(doc: Document) -> dict:
    info = dict(
        code="", name_vi="", name_en="", credits=None,
        class_hours=None, self_study_hours=None, level_vi="",
        decision_no="", faculty="", faculty_address="", prerequisites=[],
    )
    for text in _all_lines(doc):
        if not info["code"]:
            m = _CODE_RE.search(text)
            if m and len(m.group(1)) >= 4 and re.search(r'\d', m.group(1)):
                info["code"] = m.group(1).strip().upper()
            else:
                m2 = _CODE_FIELD_RE.search(text)
                if m2:
                    info["code"] = m2.group(1).strip().upper()
        if not info["name_vi"]:
            m = _VI_NAME_RE.search(text)
            if m:
                v = _clean(m.group(1))
                if 3 < len(v) < 120: info["name_vi"] = v
        if not info["name_en"]:
            m = _EN_NAME_RE.search(text)
            if m:
                v = _clean(m.group(1))
                if 3 < len(v) < 150: info["name_en"] = v
        if info["credits"] is None:
            # FIX: thử word-form trước ("three-credit course")
            wm = _CREDITS_WORD_RE.search(text)
            if wm:
                info["credits"] = _WORD_CREDITS.get(wm.group(1).lower())
            if info["credits"] is None:
                m = _CREDITS_RE.search(text)
                if m:
                    raw = m.group(1) or m.group(2)  # group1=keyword-style, group2=NNtc style
                    v = int(raw)
                    if 1 <= v <= 15: info["credits"] = v
        if info["class_hours"] is None:
            m = _CLASS_H_RE.search(text)
            if m:
                v = int(m.group(1))
                if 1 <= v <= 500: info["class_hours"] = v
        if info["self_study_hours"] is None:
            m = _SELF_H_RE.search(text)
            if m:
                v = int(m.group(1))
                if 1 <= v <= 500: info["self_study_hours"] = v
        if not info["level_vi"]:
            m = _LEVEL_RE.search(text)
            if m:
                v = _clean(m.group(1))
                if len(v) < 40: info["level_vi"] = v
        if not info["decision_no"]:
            m = _DECISION_RE.search(text)
            if m:
                v = _clean(m.group(1))
                if len(v) < 80 and re.search(r"\d", v): info["decision_no"] = v
        if not info["faculty"]:
            m = _FACULTY_RE.search(text)
            if m:
                v = _clean(m.group(1))
                if len(v) < 80 and not re.search(r"phòng|nhà a|floor|room", v, re.I):
                    info["faculty"] = v
        if not info["faculty_address"]:
            m = _ADDRESS_RE.search(text)
            if m:
                v = _clean(m.group(1))
                if len(v) < 150: info["faculty_address"] = v

    collecting = False
    for para, text in _iter_para(doc):
        if _SEC_DESC_RE.search(text) and len(text) < 70: break
        if _PREREQ_HDR.search(text):
            collecting = True
            after = re.sub(r"^.*?:\s*", "", text, count=1).strip()
            after = re.sub(r"^\*\s*", "", after).strip()
            if after and after.lower() not in ("none", "không", "không có", ""):
                info["prerequisites"].append(after)
            continue
        if collecting:
            if _PREREQ_STP.search(text) or re.match(r"^\d+\.\s", text):
                collecting = False; continue
            cleaned = re.sub(r"^[\*\+\-]\s+", "", text).strip()
            if len(cleaned) < 120 and cleaned.lower() not in ("none", "không", "không có", "-", ""):
                info["prerequisites"].append(cleaned)

    if info["credits"] is None:
        info["credits"] = _parse_credits_from_table(doc)

    if info["class_hours"] is None or info["self_study_hours"] is None:
        fb_ch, fb_sh = _parse_hours_from_table(doc)
        if info["class_hours"] is None:
            info["class_hours"] = fb_ch
        if info["self_study_hours"] is None:
            info["self_study_hours"] = fb_sh

    return info


def _parse_credits_from_table(doc: Document) -> Optional[int]:
    """Fallback: tìm credits trong bảng 2-cột.
    Hỗ trợ cả word-form, NNtc, và multiline cell kiểu:
      Label: '- Số tín chỉ:\\n+ Số tiết trên lớp:\\n+ Số giờ tự học:'
      Value: '03\\n45\\n105'
    """
    for table in doc.tables[:8]:
        for row in table.rows:
            cells = [_clean(c.text) for c in row.cells]
            unique = list(dict.fromkeys(cells))

            # FIX: multiline cell — tách label cell thành các dòng, tìm dòng "số tín chỉ"
            # rồi lấy dòng tương ứng trong value cell
            raw_cells = [c.text for c in row.cells]
            if len(raw_cells) >= 2:
                label_raw = raw_cells[0]
                value_raw = raw_cells[1] if len(set(raw_cells)) > 1 else None
                if value_raw and re.search(r"số\s*tín\s*chỉ|credits?\s*:?", label_raw, re.I):
                    label_lines = [l.strip() for l in label_raw.splitlines()]
                    value_lines = [l.strip() for l in value_raw.splitlines()]
                    for li, ll in enumerate(label_lines):
                        if re.search(r"số\s*tín\s*chỉ|credits?\s*:?$", ll, re.I):
                            val = value_lines[li] if li < len(value_lines) else value_lines[0] if value_lines else ""
                            # thử NNtc, word-form, số thuần
                            m = re.search(r"(\d+)\s*TC(?!\w)", val, re.I)
                            if not m: m = re.search(r"\b(\d+)\b", val)
                            if m:
                                v = int(m.group(1))
                                if 1 <= v <= 15: return v

            for i, cell in enumerate(unique):
                cell_stripped = re.sub(r"^[-•*+\s]+", "", cell)
                if re.search(r"số\s*tín\s*chỉ|^credits?\s*:?$|^units?\s*:?$", cell_stripped, re.I) and len(cell_stripped) < 30:
                    for j, other in enumerate(unique):
                        if j != i:
                            wm = _CREDITS_WORD_RE.search(other)
                            if wm:
                                v = _WORD_CREDITS.get(wm.group(1).lower())
                                if v: return v
                            # FIX: NNtc format
                            m = re.search(r"(\d+)\s*TC(?!\w)", other, re.I)
                            if not m: m = re.search(r"\b(\d+)\b", other)
                            if m:
                                v = int(m.group(1))
                                if 1 <= v <= 15: return v
    return None


def _parse_hours_from_table(doc: Document) -> tuple:
    _CH_KW = re.compile(
        r"số\s*giờ\s*trên\s*lớp|số\s*tiết\s*trên\s*lớp|số\s*tiết\s*lý\s*thuyết"
        r"|classroom\s*hours?|class\s*hours?|theoretical\s*hours?", re.I
    )
    _SH_KW = re.compile(r"số\s*giờ\s*tự\s*học|số\s*tiết\s*tự\s*học|self.?study\s*hours?", re.I)
    ch = sh = None
    for table in doc.tables[:10]:
        if not table.rows: continue
        for row in table.rows:
            # FIX: xử lý multiline cell: label='Số tín chỉ:\n+ Số tiết trên lớp:\n+ Số giờ tự học:'
            #      value='03\n45\n105' — ghép theo dòng
            raw_cells = [c.text for c in row.cells]
            if len(raw_cells) >= 2 and len(set(raw_cells)) > 1:
                label_raw = raw_cells[0]
                value_raw = raw_cells[1]
                if "\n" in label_raw:
                    label_lines = [l.strip() for l in label_raw.splitlines()]
                    value_lines = [l.strip() for l in value_raw.splitlines()]
                    for li, ll in enumerate(label_lines):
                        val = value_lines[li] if li < len(value_lines) else ""
                        m_v = re.search(r"\b(\d{2,3})\b", val)
                        if m_v:
                            v = int(m_v.group(1))
                            if ch is None and _CH_KW.search(ll) and 1 <= v <= 500:
                                ch = v
                            if sh is None and _SH_KW.search(ll) and 1 <= v <= 500:
                                sh = v

            cells = [_clean(c.text) for c in row.cells]
            unique = list(dict.fromkeys(cells))
            if len(unique) < 2: continue
            label, rest = unique[0], " ".join(unique[1:])
            if ch is None and _CH_KW.search(label):
                m = re.search(r"\b(\d{2,3})\b", rest)
                if m:
                    v = int(m.group(1))
                    if 1 <= v <= 500: ch = v
            if sh is None and _SH_KW.search(label):
                m = re.search(r"\b(\d{2,3})\b", rest)
                if m:
                    v = int(m.group(1))
                    if 1 <= v <= 500: sh = v
        if ch is not None and sh is not None: break
    return ch, sh


# ── Section 2: giảng viên ─────────────────────────────────────────────────────

_INSTR_LINE_RE = re.compile(r"^(\d+)\.\s*(.+)", re.I)
_INSTR_ROLE_RE = re.compile(
    r"^((?:PhD\.?|MSc\.?|MA\.?|Dr\.?|Assoc\.?\s*Prof\.?|Prof\.?|TS\.?|ThS\.?|PGS\.?|GS\.?)"
    r"(?:\.?\s*(?:PhD|TS)\.?)?)?\s*",
    re.I,
)
_INSTR_HDR_RE = re.compile(
    r"(?:lecturers?|course\s*instructors?|giảng\s*viên(?:\s*giảng\s*dạy)?)\s*:?\s*$",
    re.I,
)

def _parse_instructors(doc: Document) -> list:
    instructors = []
    for table in doc.tables[:5]:
        if not table.rows: continue
        hdr = " ".join(_clean(c.text).lower() for c in table.rows[0].cells)
        # FIX: "họ tên" (không có "và") + "họ và tên"
        if not re.search(r"họ\s*(?:và\s*)?tên|full\s*name|lecturer", hdr): continue

        name_col = email_col = -1
        for i, h in enumerate(_clean(c.text).lower() for c in table.rows[0].cells):
            if re.search(r"họ\s*(?:và\s*)?tên|full\s*name|^name$", h): name_col  = i
            if re.search(r"email|e-mail",                             h): email_col = i

        if name_col < 0: continue
        order = 1
        for row in table.rows[1:]:
            cells = [_clean(c.text) for c in row.cells]
            if not cells or _is_placeholder_row(cells): continue
            if name_col >= len(cells): continue
            raw_name = cells[name_col]
            if not raw_name or len(raw_name) < 3: continue

            role_m = re.match(r"^((?:GS|PGS|TS|ThS|CN)\.?\s*(?:TS\.?\s*)?)", raw_name, re.I)
            role_text = role_m.group(1).strip() if role_m else None
            full_name = raw_name[role_m.end():].strip() if role_m else raw_name

            email = None
            if email_col >= 0 and email_col < len(cells):
                em = _EMAIL_RE.search(cells[email_col])
                if em: email = em.group(0)
            # FIX: tìm email trong toàn bộ row nếu chưa có
            if not email:
                for cell_text in cells:
                    em = _EMAIL_RE.search(cell_text)
                    if em: email = em.group(0); break

            if full_name:
                instructors.append(Instructor(
                    full_name=full_name, email=email,
                    role_text=role_text, order_no=order,
                ))
                order += 1
        if instructors: break

    # Fallback: dạng paragraph "1. PhD. Name   Tel: ...   Email: ..."
    if not instructors:
        collecting = False
        order = 1
        for para, text in _iter_para(doc):
            if _INSTR_HDR_RE.search(text) and len(text) < 80:
                collecting = True
                continue
            if collecting:
                if re.match(r"^[*+]\s*(?:kind|course\s*lang|meeting)", text, re.I):
                    break
                lm = _INSTR_LINE_RE.match(text)
                if not lm:
                    if re.match(r"^\d+\.\s+[A-Z\*]", text):
                        break
                    continue
                raw = lm.group(2).strip()
                rm = _INSTR_ROLE_RE.match(raw)
                role = rm.group(1).strip() if rm and rm.group(1) else None
                rest = raw[rm.end():].strip() if rm else raw
                name = re.split(r"\s{2,}|\s+(?:Tel|Email|Phone)\s*:", rest, maxsplit=1)[0].strip()
                # Loại bỏ title suffix: ", PhD", ", Dr"
                name = re.sub(r",\s*(?:PhD|Dr\.?|MSc|MA)\.?\s*$", "", name, flags=re.I).strip()
                email_m = _EMAIL_RE.search(text)
                email = email_m.group(0) if email_m else None
                if name and len(name) >= 3:
                    instructors.append(Instructor(
                        full_name=name, email=email,
                        role_text=role, order_no=order,
                    ))
                    order += 1

    return instructors


# ── Section 4: tài liệu học tập ───────────────────────────────────────────────

def _parse_resources(doc: Document) -> list:
    resources = []
    collecting = False
    current_group = "other"
    order = 1
    GROUPS = [
        (re.compile(r"giáo\s*trình|text\s*book|required\s*textbook|course\s*book",   re.I), "textbook"),
        (re.compile(r"tài\s*liệu\s*tham\s*khảo|additional\s*ref|reference|recommended\s*texts?|other\s*readings?", re.I), "reference"),
        (re.compile(r"phần\s*mềm|software",                       re.I), "software"),
    ]
    STOP = re.compile(
        r"(?:\d+\.\s*)?(?:mục\s*tiêu\s*học\s*phần|course\s*goals?|"
        r"chuẩn\s*đầu\s*ra|đánh\s*giá|assessment\b|grading|kế\s*hoạch|"
        r"quy\s*định|course\s*requirements?|lesson\s*content|course\s*content)",
        re.I,
    )

    for para, text in _iter_para(doc):
        if _SEC_RES_RE.search(text) and len(text) < 120:
            collecting = True; continue
        if collecting:
            if STOP.search(text) and len(text) < 80:
                break

            # Kiểm tra sub-group heading (giáo trình / tài liệu tham khảo / phần mềm)
            new_group = None
            for pat, gname in GROUPS:
                if pat.search(text) and len(text) < 60:
                    new_group = gname; break
            if new_group:
                current_group = new_group; continue

            # FIX: chấp nhận plain text (không bắt buộc bullet/số)
            # Bỏ qua ghi chú trong ngoặc: "(Learning resources: ...)"
            if text.startswith("(") and text.endswith(")"):
                continue
            # Bỏ qua sub-heading ngắn có số: "4.1. Giáo trình", "4.2. Tài liệu tham khảo"
            if re.match(r"^\d+\.\d+\.?\s+\w", text) and len(text) < 60:
                continue
            # Chỉ lấy text đủ dài (tránh heading/label ngắn)
            if len(text) > 20:
                citation = re.sub(r"^[\[\]\-\*•\d\.]+\s*", "", text).strip()
                if citation:
                    resources.append(Resource(
                        resource_group=current_group,
                        citation_text=citation,
                        order_no=order,
                    ))
                    order += 1
    return resources


# ── Section 5: mục tiêu học phần ─────────────────────────────────────────────

def _parse_goals(doc: Document) -> list:
    goals = []
    for table in doc.tables:
        if not table.rows: continue
        hdr = " ".join(_clean(c.text).lower() for c in table.rows[0].cells)
        if not re.search(r"mục\s*tiêu|goal", hdr): continue
        if re.search(r"\bclo\b", hdr): continue

        code_col = desc_col = -1
        for i, h in enumerate(_clean(c.text).lower() for c in table.rows[0].cells):
            # FIX: thêm "^tt$" và "^stt$" — cột thứ tự chứa mã G1/G2/CG1/CG2
            if re.search(r"mã|code|^goals?$|^tt$|^stt$", h): code_col = i
            if re.search(r"mô\s*tả|description|goals?\s*desc", h) and i != code_col: desc_col = i

        # FIX: Nếu không tìm được code_col qua header → data-driven scan
        if code_col < 0 and desc_col >= 0:
            for row in table.rows[1:5]:
                cells = [_clean(c.text) for c in row.cells]
                if _is_placeholder_row(cells): continue
                for i, cell in enumerate(cells):
                    if i != desc_col and _GOAL_RE.match(cell):
                        code_col = i
                        break
                if code_col >= 0: break

        # Fallback cuối: nếu vẫn không tìm được, dùng cột đầu tiên khác desc_col
        if code_col < 0 and desc_col > 0:
            code_col = 0

        if code_col < 0 or desc_col < 0: continue

        order = 1
        for row in table.rows[1:]:
            cells = [_clean(c.text) for c in row.cells]
            if not cells or _is_placeholder_row(cells): continue
            if code_col >= len(cells) or desc_col >= len(cells): continue
            code = cells[code_col]
            desc = cells[desc_col]
            if not code or not desc or len(desc) < 5: continue
            m = _GOAL_RE.match(code)
            if not m: continue
            goals.append(Goal(
                goal_code=m.group(1).upper(),
                description_vi=desc,
                display_order=order,
            ))
            order += 1
        if goals: break
    # Fallback: EN-format goals as numbered paragraphs "1) ... N)"
    if not goals:
        goals = _parse_goals_from_paragraphs(doc)
    return goals


def _parse_goals_from_paragraphs(doc: Document) -> list:
    """Fallback: EN-format goals listed as '1) ...' paragraphs under SLOs/Objectives header."""
    _GOAL_HDR = re.compile(
        r"student\s*learning\s*outcomes?|main\s*course\s*(?:learning\s*)?objectives?"
        r"|course\s*objectives?|by\s*the\s*end\s*of\s*this\s*course",
        re.I,
    )
    _GOAL_PARA_RE = re.compile(r"^(\d+)[)\.]\s*(.+)", re.I)
    _STOP = re.compile(
        r"^\d+\.\s+[A-Z]|course\s*content|tentative\s*schedule"
        r"|required\s*textbook|assessment|grading",
        re.I,
    )
    goals, collecting, order = [], False, 1
    for para, text in _iter_para(doc):
        if _GOAL_HDR.search(text) and len(text) < 100:
            collecting = True; continue
        if collecting:
            if (_STOP.search(text) and len(text) < 80) or (re.match(r"^\d+\.\s", text) and len(text) < 60):
                break
            m = _GOAL_PARA_RE.match(text)
            if m:
                desc = m.group(2).rstrip(",").strip()
                if len(desc) > 10:
                    goals.append(Goal(
                        goal_code=f"G{order}", description_vi=desc, display_order=order,
                    ))
                    order += 1
    return goals


# ── Section 3: mô tả học phần ────────────────────────────────────────────────

def _parse_description(doc: Document) -> Optional[str]:
    collecting, chunks = False, []
    for para, text in _iter_para(doc):
        if _SEC_DESC_RE.search(text) and len(text) < 70:
            collecting = True; continue
        if collecting:
            style = (para.style.name or "").lower()
            if ("heading" in style and text and len(text) < 100) or (re.match(r"^\d+\.\s", text) and len(text) < 80):
                break
            if re.match(r"^nội\s*dung\s*chính|^main\s*content", text, re.I): break
            chunks.append(text)
            if len(chunks) >= 5: break
    if chunks: return " ".join(chunks)
    for tbl in doc.tables[:3]:
        for row in tbl.rows:
            for cell in row.cells:
                t = _clean(cell.text)
                if len(t) > 120 and not _CLO_CODE_RE.search(t): return t[:1000]
    return None


# ── Section 6: CLO ────────────────────────────────────────────────────────────

def _is_clo_master(table: Table) -> bool:
    if len(table.rows) < 3: return False
    hdr = " ".join(_clean(c.text).lower() for c in table.rows[0].cells)
    return (
        re.search(r"\bclo", hdr) and
        re.search(r"mô\s*tả|description\*?|outcome", hdr) and
        (re.search(r"mục\s*tiêu|goals?\b", hdr) or
         re.search(r"mức\s*độ|level\*?|proficiency|teaching", hdr))
    )

def _extract_clos(table: Table) -> list:
    if not table.rows: return []
    hdr = [_clean(c.text).lower() for c in table.rows[0].cells]
    goal_col = desc_col = code_col = level_col = -1
    for i, h in enumerate(hdr):
        if re.search(r"mô\s*tả|description\*?|diễn\s*giải|outcome", h): desc_col = i
    for i, h in enumerate(hdr):
        if re.search(r"mục\s*tiêu|goals?\b",                h): goal_col  = i
        if re.search(r"mức\s*độ|level\*?|proficiency|teaching", h): level_col = i
        if re.search(r"\bclo", h) and i != desc_col:         code_col  = i
    if code_col < 0 or desc_col < 0: return []

    # FIX: phát hiện kiểu mã số thuần (không có tiền tố CLO) — vd "1.1", "2.3"
    # Scan vài dòng đầu để xem cột code_col có dùng dạng X.Y không
    _NUMERIC_CLO_RE = re.compile(r"^(\d+)\.(\d+)$")
    use_numeric_clo = False
    for row in table.rows[2:6]:
        cells = [_clean(c.text) for c in row.cells]
        if code_col < len(cells) and _NUMERIC_CLO_RE.match(cells[code_col]):
            use_numeric_clo = True
            break

    clos, cur_goal, order = [], "", 1
    for row in table.rows[1:]:
        cells = [_clean(c.text) for c in row.cells]
        if not cells or _is_placeholder_row(cells): continue
        if goal_col >= 0 and goal_col < len(cells) and cells[goal_col]:
            m = _GOAL_RE.match(cells[goal_col])
            if m: cur_goal = m.group(1).upper()
        if code_col >= len(cells): continue
        raw_code = cells[code_col]

        if use_numeric_clo:
            # FIX: dạng "1.1" → CLO1.1
            nm = _NUMERIC_CLO_RE.match(raw_code)
            if nm:
                code = f"CLO{nm.group(1)}.{nm.group(2)}"
            else:
                m = _CLO_CODE_RE.search(raw_code)
                if not m: continue
                code = _norm_clo(m.group(1), m.group(2))
        else:
            m = _CLO_CODE_RE.search(raw_code)
            if not m: continue
            code = _norm_clo(m.group(1), m.group(2))

        if desc_col >= len(cells): continue
        desc = cells[desc_col]
        if not desc or len(desc) < 5: continue
        level = None
        if level_col >= 0 and level_col < len(cells):
            lv = cells[level_col]
            if lv and not _is_placeholder_row([lv]):
                level = _roman_to_float(lv)
        clos.append(CLO(
            clo_code=code, description_vi=desc,
            goal_code=cur_goal or None, attainment_level=level,
            display_order=order,
        ))
        order += 1
    return clos

def _parse_clos(doc: Document) -> list:
    for tbl in doc.tables:
        if _is_clo_master(tbl):
            result = _extract_clos(tbl)
            if result: return result
    return []


# ── Section 7: đánh giá học phần ─────────────────────────────────────────────

def _parse_assessments(doc: Document) -> list:
    assessments = []
    for table in doc.tables:
        if not table.rows: continue
        hdr = " ".join(_clean(c.text).lower() for c in table.rows[0].cells)
        if not re.search(r"hình\s*thức\s*đánh\s*giá|assessment|component", hdr): continue
        if not re.search(r"tỷ\s*lệ|weight|%", hdr): continue

        kind_col = content_col = time_col = clo_col = tool_col = pct_col = -1
        for i, h in enumerate(_clean(c.text).lower() for c in table.rows[0].cells):
            if re.search(r"hình\s*thức|assessment|type|component",        h): kind_col    = i
            if re.search(r"nội\s*dung|bài\s*đánh\s*giá|content", h) and i != kind_col: content_col = i
            if re.search(r"thời\s*điểm|thời\s*gian|timing|week",   h): time_col    = i
            if re.search(r"\bclo\b|\bclos\b",                    h): clo_col     = i
            if re.search(r"công\s*cụ|rubric|tool|tiêu\s*chí",   h): tool_col    = i
            if re.search(r"tỷ\s*lệ|weight|%",                   h): pct_col     = i

        if kind_col < 0 and pct_col < 0: continue

        order = 1
        for row in table.rows[1:]:
            cells = [_clean(c.text) for c in row.cells]
            if not cells or _is_placeholder_row(cells): continue
            title_raw = cells[kind_col] if kind_col >= 0 and kind_col < len(cells) else ""
            if not title_raw: continue

            pct = None
            if pct_col >= 0 and pct_col < len(cells):
                pm = re.search(r"(\d+(?:[.,]\d+)?)\s*%?", cells[pct_col])
                if pm:
                    try: pct = float(pm.group(1).replace(",", "."))
                    except: pass

            clo_text  = cells[clo_col] if clo_col >= 0 and clo_col < len(cells) else ""
            clo_codes = _extract_clo_codes(clo_text)

            assessments.append(Assessment(
                assessment_kind = _guess_assessment_kind(title_raw),
                title           = title_raw,
                weight_percent  = pct,
                timing_text     = cells[time_col]    if time_col    >= 0 and time_col    < len(cells) else None,
                content_text    = cells[content_col] if content_col >= 0 and content_col < len(cells) else None,
                evaluation_tool = cells[tool_col]    if tool_col    >= 0 and tool_col    < len(cells) else None,
                clo_codes       = clo_codes,
                display_order   = order,
            ))
            order += 1
        if assessments: break
    return assessments


# ── Section 8: kế hoạch dạy học ──────────────────────────────────────────────

def _parse_weeks(doc: Document) -> list:
    weeks = []
    for table in doc.tables:
        if not table.rows: continue
        hdr = " ".join(_clean(c.text).lower() for c in table.rows[0].cells)

        # Case A: standard table with tuần/week + nội dung/content keywords
        has_week_kw    = bool(re.search(r"tuần|week", hdr))
        has_content_kw = bool(re.search(r"nội\s*dung|content|topic", hdr))

        # Case B: "No | Contents" table used in some EN-format docs (TIKT style)
        first_h = _clean(table.rows[0].cells[0].text).lower()
        is_no_contents = (
            bool(re.match(r"^no\.?$|^#$|^stt\.?$", first_h)) and
            bool(re.search(r"content|topic", hdr))
        )

        if not (has_week_kw and has_content_kw) and not is_no_contents:
            continue

        week_col = content_col = reading_col = activity_col = assess_col = clo_col = -1

        if is_no_contents and not (has_week_kw and has_content_kw):
            # Case B: No→week_col=0, find content col
            week_col = 0
            for i, h in enumerate(_clean(c.text).lower() for c in table.rows[0].cells):
                if i == 0: continue
                if re.search(r"content|topic", h)    and content_col  < 0: content_col  = i
                if re.search(r"tài\s*liệu|reading|reference", h):          reading_col  = i
                if re.search(r"hoạt\s*động|activit|teaching", h):          activity_col = i
                if re.search(r"đánh\s*giá|assessment(?!\s*tool)", h):      assess_col   = i
                if re.search(r"\bclo\b|\bclos\b|rcx?\d*", h):              clo_col      = i
        else:
            # Case A: find columns by header keywords
            for i, h in enumerate(_clean(c.text).lower() for c in table.rows[0].cells):
                # FIX: use ^week (not ^week$) to match "week/class", "weeks"
                if re.search(r"^tuần\b|^week", h):                          week_col     = i
                if re.search(r"nội\s*dung|content|topic", h) and i != week_col: content_col = i
                if re.search(r"tài\s*liệu|reading|reference", h):          reading_col  = i
                if re.search(r"hoạt\s*động|activit|teaching", h):          activity_col = i
                if re.search(r"đánh\s*giá|assessment(?!\s*tool)|rubric", h): assess_col  = i
                if re.search(r"\bclo\b|\bclos\b",                  h):      clo_col      = i

        if week_col < 0 or content_col < 0: continue

        seq_no = 0
        for row in table.rows[1:]:
            cells = [_clean(c.text) for c in row.cells]
            if not cells or _is_placeholder_row(cells): continue
            if week_col >= len(cells): continue

            week_raw    = cells[week_col]
            content_raw = cells[content_col] if content_col < len(cells) else ""
            if not content_raw: continue

            # Skip secondary-header rows (e.g. "No | Contents | Theory | Practice")
            if (re.match(r"^no\.?$|^#$|^tuần$|^week$", week_raw, re.I) and
                    re.search(r"content|topic|nội\s*dung", content_raw, re.I) and len(content_raw) < 30):
                continue

            wm = re.search(r"(\d+)", week_raw) if week_raw else None
            if wm:
                week_no = int(wm.group(1)); seq_no = week_no
            else:
                cm = re.match(r"(?:Lecture|Seminar|Lab|Week)\s*(\d+)", content_raw, re.I)
                if cm:
                    week_no = int(cm.group(1)); seq_no = week_no
                else:
                    seq_no += 1; week_no = seq_no

            session_m     = re.match(r"((?:Lecture|Seminar|Lab|Ôn tập|Kiểm tra)\s*\d*)", content_raw, re.I)
            session_label = session_m.group(1).strip() if session_m else None

            clo_text  = cells[clo_col] if clo_col >= 0 and clo_col < len(cells) else ""
            clo_codes = _extract_clo_codes(clo_text)

            weeks.append(Week(
                week_no       = week_no,
                session_label = session_label,
                content_text  = content_raw,
                reading_text  = cells[reading_col]  if reading_col  >= 0 and reading_col  < len(cells) else None,
                teaching_learning_activities = cells[activity_col] if activity_col >= 0 and activity_col < len(cells) else None,
                assessment_text = cells[assess_col] if assess_col  >= 0 and assess_col  < len(cells) else None,
                clo_codes     = clo_codes,
            ))
        if weeks: break
    return weeks


# ── Section 9: quy định học phần ─────────────────────────────────────────────

def _parse_requirements(doc: Document) -> list:
    requirements = []
    order = 1
    collecting = False
    STOP = re.compile(
        r"(?:\d+\.\s*)?(?:phụ\s*lục|rubric|ma\s*trận|tài\s*liệu\s*tham\s*khảo)", re.I
    )
    for para, text in _iter_para(doc):
        if _SEC_REQ_RE.search(text) and len(text) < 80:
            collecting = True; continue
        if collecting:
            if STOP.search(text) and len(text) < 60: break
            if re.match(r"^\d+\.\s", text) and len(text) < 60 and not re.search(
                r"[a-záàảãạăắặằẵặêếệềểẽôốộồổỗơớợờởỡưứựừửữ]", text, re.I
            ):
                break
            cleaned = re.sub(r"^[\*\+\-•]\s*", "", text).strip()
            if len(cleaned) > 10:
                requirements.append(Requirement(
                    requirement_group=_guess_req_group(cleaned),
                    detail_text=cleaned,
                    order_no=order,
                ))
                order += 1

    if not requirements:
        for table in doc.tables:
            if not table.rows: continue
            hdr = " ".join(_clean(c.text).lower() for c in table.rows[0].cells)
            if not re.search(r"quy\s*định|requirement|điều\s*kiện", hdr): continue
            for row in table.rows[1:]:
                cells = [_clean(c.text) for c in row.cells]
                if not cells or _is_placeholder_row(cells): continue
                text = " ".join(c for c in cells if c)
                if len(text) > 10:
                    requirements.append(Requirement(
                        requirement_group=_guess_req_group(text),
                        detail_text=text,
                        order_no=order,
                    ))
                    order += 1
            if requirements: break

    return requirements


# ── Phụ lục: rubrics ─────────────────────────────────────────────────────────

def _parse_rubrics(doc: Document) -> list:
    rubrics = []
    rubric_order = 1
    for table in doc.tables:
        if not table.rows or len(table.rows) < 3: continue
        hdr = [_clean(c.text).lower() for c in table.rows[0].cells]
        hdr_str = " ".join(hdr)
        if not re.search(r"tiêu\s*chí|criterion|rubric", hdr_str, re.I): continue
        if not re.search(r"xuất\s*sắc|excellent|tốt|good|trung\s*bình|average|yếu|poor", hdr_str, re.I): continue

        name_col = clo_col = weight_col = exc_col = good_col = avg_col = poor_col = -1
        for i, h in enumerate(hdr):
            if re.search(r"tiêu\s*chí|criterion",                  h): name_col   = i
            if re.search(r"\bclo\b",                                h): clo_col    = i
            if re.search(r"trọng\s*số|weight|tỷ\s*lệ",            h): weight_col = i
            if re.search(r"xuất\s*sắc|excellent",                  h): exc_col    = i
            if re.search(r"tốt|good", h) and i != exc_col:            good_col  = i
            if re.search(r"trung\s*bình|average|tb",                h): avg_col    = i
            if re.search(r"yếu|poor|không\s*đạt",                  h): poor_col   = i

        if name_col < 0: continue

        rubric_title = f"Rubric {rubric_order}"
        for para in doc.paragraphs:
            t = _clean(para.text)
            if re.search(r"rubric\s*\d+|phụ\s*lục\s*\d+", t, re.I) and len(t) < 80:
                rubric_title = t

        rubric_weight = None
        for h in hdr:
            m = re.search(r"(\d+(?:[.,]\d+)?)\s*%", h)
            if m:
                try: rubric_weight = float(m.group(1).replace(",", ".")); break
                except: pass

        criteria = []
        crit_order = 1
        for row in table.rows[1:]:
            cells = [_clean(c.text) for c in row.cells]
            if not cells or _is_placeholder_row(cells): continue
            if name_col >= len(cells): continue
            name = cells[name_col]
            if not name or len(name) < 2: continue

            weight = None
            if weight_col >= 0 and weight_col < len(cells):
                wm = re.search(r"(\d+(?:[.,]\d+)?)\s*%?", cells[weight_col])
                if wm:
                    try: weight = float(wm.group(1).replace(",", "."))
                    except: pass

            criteria.append(RubricCriterion(
                criterion_name = name,
                clo_text       = cells[clo_col]  if clo_col  >= 0 and clo_col  < len(cells) else None,
                weight_percent = weight,
                excellent_desc = cells[exc_col]  if exc_col  >= 0 and exc_col  < len(cells) else None,
                good_desc      = cells[good_col] if good_col >= 0 and good_col < len(cells) else None,
                average_desc   = cells[avg_col]  if avg_col  >= 0 and avg_col  < len(cells) else None,
                poor_desc      = cells[poor_col] if poor_col >= 0 and poor_col < len(cells) else None,
                display_order  = crit_order,
            ))
            crit_order += 1

        if criteria:
            rubrics.append(Rubric(
                title         = rubric_title,
                weight_percent= rubric_weight,
                criteria      = criteria,
                display_order = rubric_order,
            ))
            rubric_order += 1

    return rubrics


# ── Detect flags ──────────────────────────────────────────────────────────────

def _detect_is_thuc_te(name_vi: str, name_en: str, filename: str) -> bool:
    """True nếu học phần là Chuyên đề thực tế (field trip) — không có kế hoạch tuần."""
    combined = f"{name_vi} {name_en} {filename}"
    return bool(re.search(
        r"chuyên\s*đề\s*thực\s*tế|practical\s*training|field\s*trip", combined, re.I
    ))

def _detect_has_clo_table(doc: Document) -> bool:
    """True nếu tài liệu có bảng CLO chuẩn (không phải chỉ có dạng paragraph)."""
    return any(_is_clo_master(tbl) for tbl in doc.tables)


# ── Main entry ────────────────────────────────────────────────────────────────

def parse_docx(file_bytes: bytes, filename: str) -> SyllabusRecord:
    doc             = Document(BytesIO(file_bytes))
    fn_code, fn_name = parse_filename(filename)
    info            = _parse_general_info(doc)
    description     = _parse_description(doc)
    clos            = _parse_clos(doc)

    name_vi = info["name_vi"] or fn_name.replace("_", " ")
    name_en = info["name_en"] or None
    if not name_vi and name_en: name_vi = name_en

    prereq_text = "; ".join(info["prerequisites"]) if info["prerequisites"] else None
    course_code = info["code"] or fn_code

    is_thuc_te    = _detect_is_thuc_te(name_vi or "", name_en or "", filename)
    has_clo_table = _detect_has_clo_table(doc)

    return SyllabusRecord(
        course_code_snapshot    = course_code,
        course_name_vi_snapshot = name_vi,
        course_name_en_snapshot = name_en,
        credits                 = info["credits"],
        class_hours             = info["class_hours"],
        self_study_hours        = info["self_study_hours"],
        level_name_vi           = info["level_vi"] or None,
        decision_no             = info["decision_no"] or None,
        managing_faculty_text   = info["faculty"] or None,
        faculty_address_text    = info["faculty_address"] or None,
        prerequisite_text       = prereq_text,
        course_description      = description,
        status                  = "draft",
        instructors             = _parse_instructors(doc),
        resources               = _parse_resources(doc),
        goals                   = _parse_goals(doc),
        clos                    = clos,
        assessments             = _parse_assessments(doc),
        weeks                   = _parse_weeks(doc),
        requirements            = _parse_requirements(doc),
        rubrics                 = _parse_rubrics(doc),
        source_file             = filename,
        is_thuc_te              = is_thuc_te,
        has_clo_table           = has_clo_table,
    )
